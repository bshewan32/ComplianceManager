"""
Enhanced Document Scanner with Content-Based Matching
Improves on the basic filename matching with full text extraction and fuzzy matching.
"""

import os
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from collections import defaultdict

# File processing imports
import PyPDF2
import docx
from openpyxl import load_workbook
from PIL import Image

# Fuzzy matching
from difflib import SequenceMatcher


class EnhancedDocumentScanner:
    """
    Advanced document scanner that extracts full text content and uses
    intelligent matching algorithms to map documents to compliance clauses.
    """
    
    def __init__(self):
        self.supported_extensions = {
            '.pdf': self._extract_pdf_text,
            '.docx': self._extract_docx_text,
            '.doc': self._extract_docx_text,
            '.xlsx': self._extract_xlsx_text,
            '.xls': self._extract_xlsx_text,
            '.txt': self._extract_txt_text,
        }
        
    # ==================== TEXT EXTRACTION ====================
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF files."""
        try:
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            return ' '.join(text)
        except Exception as e:
            print(f"Error extracting PDF text from {file_path}: {e}")
            return ""
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from Word documents."""
        try:
            doc = docx.Document(file_path)
            text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return ' '.join(text)
        except Exception as e:
            print(f"Error extracting DOCX text from {file_path}: {e}")
            return ""
    
    def _extract_xlsx_text(self, file_path: str) -> str:
        """Extract text from Excel spreadsheets."""
        try:
            workbook = load_workbook(file_path, read_only=True, data_only=True)
            text = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text.append(f"Sheet: {sheet_name}")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = ' '.join([str(cell) for cell in row if cell is not None])
                    if row_text.strip():
                        text.append(row_text)
            
            return ' '.join(text)
        except Exception as e:
            print(f"Error extracting XLSX text from {file_path}: {e}")
            return ""
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting TXT text from {file_path}: {e}")
            return ""
    
    def extract_text_from_file(self, file_path: str) -> Tuple[str, str]:
        """
        Extract text content from a file.
        
        Returns:
            Tuple of (full_text, filename_text)
        """
        file_ext = Path(file_path).suffix.lower()
        filename = Path(file_path).stem  # Filename without extension
        
        # Extract full text content
        if file_ext in self.supported_extensions:
            full_text = self.supported_extensions[file_ext](file_path)
        else:
            full_text = ""
        
        return full_text, filename
    
    # ==================== KEYWORD EXTRACTION ====================
    
    def extract_keywords(self, text: str, top_n: int = 50) -> List[str]:
        """
        Extract meaningful keywords from text.
        Filters out common stop words and keeps domain-relevant terms.
        """
        # Convert to lowercase
        text = text.lower()
        
        # Remove special characters but keep spaces
        text = re.sub(r'[^a-z0-9\s]', ' ', text)
        
        # Split into words
        words = text.split()
        
        # Common stop words to exclude
        stop_words = {
            'the', 'be', 'to', 'of', 'and', 'a', 'in', 'that', 'have', 'i',
            'it', 'for', 'not', 'on', 'with', 'he', 'as', 'you', 'do', 'at',
            'this', 'but', 'his', 'by', 'from', 'they', 'we', 'say', 'her', 'she',
            'or', 'an', 'will', 'my', 'one', 'all', 'would', 'there', 'their',
            'what', 'so', 'up', 'out', 'if', 'about', 'who', 'get', 'which', 'go',
            'me', 'when', 'make', 'can', 'like', 'time', 'no', 'just', 'him', 'know',
            'take', 'people', 'into', 'year', 'your', 'good', 'some', 'could', 'them',
            'see', 'other', 'than', 'then', 'now', 'look', 'only', 'come', 'its', 'over',
            'also', 'back', 'after', 'use', 'two', 'how', 'our', 'work', 'first', 'well',
            'way', 'even', 'new', 'want', 'because', 'any', 'these', 'give', 'day', 'most', 'us'
        }
        
        # Filter words
        keywords = [
            word for word in words 
            if len(word) > 2 and word not in stop_words
        ]
        
        # Count frequency
        word_freq = defaultdict(int)
        for word in keywords:
            word_freq[word] += 1
        
        # Return top N most frequent
        sorted_keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, freq in sorted_keywords[:top_n]]
    
    # ==================== MATCHING ALGORITHMS ====================
    
    def fuzzy_match_score(self, text1: str, text2: str) -> float:
        """
        Calculate fuzzy match score between two strings.
        Returns a score between 0 and 1.
        """
        return SequenceMatcher(None, text1.lower(), text2.lower()).ratio()
    
    def keyword_overlap_score(self, doc_keywords: List[str], clause_keywords: List[str]) -> float:
        """
        Calculate overlap between document and clause keywords.
        Returns a score between 0 and 1.
        """
        if not doc_keywords or not clause_keywords:
            return 0.0
        
        doc_set = set(doc_keywords)
        clause_set = set(clause_keywords)
        
        intersection = len(doc_set & clause_set)
        union = len(doc_set | clause_set)
        
        return intersection / union if union > 0 else 0.0
    
    def calculate_match_score(
        self,
        doc_filename: str,
        doc_text: str,
        clause_number: str,
        clause_title: str,
        clause_description: str
    ) -> float:
        """
        Calculate comprehensive match score between document and clause.
        
        Scoring weights:
        - Clause number in filename: 40%
        - Fuzzy match with clause title: 30%
        - Keyword overlap: 30%
        """
        score = 0.0
        
        # 1. Check for clause number in filename (40% weight)
        if clause_number:
            # Clean clause number (e.g., "6.1.2" or "6.1.2.1")
            clause_num_clean = clause_number.replace('.', '')
            if clause_num_clean in doc_filename.replace('.', '').replace('-', '').replace('_', ''):
                score += 0.40
        
        # 2. Fuzzy match filename with clause title (30% weight)
        title_match = self.fuzzy_match_score(doc_filename, clause_title)
        score += title_match * 0.30
        
        # 3. Keyword overlap between doc content and clause description (30% weight)
        if doc_text and clause_description:
            doc_keywords = self.extract_keywords(doc_text, top_n=30)
            clause_keywords = self.extract_keywords(clause_description, top_n=30)
            keyword_score = self.keyword_overlap_score(doc_keywords, clause_keywords)
            score += keyword_score * 0.30
        
        return min(score, 1.0)  # Cap at 1.0
    
    def match_document_to_clauses(
        self,
        file_path: str,
        clauses: List[Dict],
        threshold: float = 0.3
    ) -> List[Tuple[int, float, str]]:
        """
        Match a document to the most relevant clauses.
        
        Args:
            file_path: Path to the document
            clauses: List of clause dictionaries with id, clause_number, title, description
            threshold: Minimum score to consider a match (default 0.3)
        
        Returns:
            List of tuples (clause_id, match_score, reason) sorted by score descending
        """
        # Extract text from document
        doc_text, doc_filename = self.extract_text_from_file(file_path)
        
        # Calculate scores for each clause
        matches = []
        for clause in clauses:
            score = self.calculate_match_score(
                doc_filename=doc_filename,
                doc_text=doc_text,
                clause_number=clause.get('clause_number', ''),
                clause_title=clause.get('title', ''),
                clause_description=clause.get('description', '')
            )
            
            if score >= threshold:
                # Generate reason for match
                reason = self._generate_match_reason(
                    score, doc_filename, clause.get('clause_number', ''), 
                    clause.get('title', '')
                )
                matches.append((clause['id'], score, reason))
        
        # Sort by score descending
        matches.sort(key=lambda x: x[1], reverse=True)
        return matches
    
    def _generate_match_reason(
        self, 
        score: float, 
        filename: str, 
        clause_number: str, 
        clause_title: str
    ) -> str:
        """Generate human-readable reason for why document matched clause."""
        reasons = []
        
        if clause_number and clause_number.replace('.', '') in filename.replace('.', '').replace('-', '').replace('_', ''):
            reasons.append(f"contains clause number '{clause_number}'")
        
        title_match = self.fuzzy_match_score(filename, clause_title)
        if title_match > 0.5:
            reasons.append(f"filename similar to clause title ({int(title_match*100)}% match)")
        
        if reasons:
            return "Matched because filename " + " and ".join(reasons)
        else:
            return f"Content keywords match clause description ({int(score*100)}% confidence)"
    
    # ==================== BATCH SCANNING ====================
    
    def scan_folder(
        self,
        folder_path: str,
        clauses: List[Dict],
        match_threshold: float = 0.3,
        max_matches_per_doc: int = 3
    ) -> Dict:
        """
        Scan an entire folder and match documents to clauses.
        
        Args:
            folder_path: Path to folder containing documents
            clauses: List of clause dictionaries
            match_threshold: Minimum score to consider a match
            max_matches_per_doc: Maximum number of clause matches per document
        
        Returns:
            Dictionary with scan results and statistics
        """
        if not os.path.exists(folder_path):
            raise ValueError(f"Folder path does not exist: {folder_path}")
        
        results = {
            'documents_scanned': 0,
            'documents_matched': 0,
            'total_matches': 0,
            'unmatched_documents': [],
            'matches': []  # List of {file_path, clause_matches: [(clause_id, score, reason)]}
        }
        
        # Scan all files recursively
        for root, dirs, files in os.walk(folder_path):
            for filename in files:
                file_path = os.path.join(root, filename)
                file_ext = Path(file_path).suffix.lower()
                
                # Skip unsupported file types
                if file_ext not in self.supported_extensions:
                    continue
                
                results['documents_scanned'] += 1
                
                # Match document to clauses
                matches = self.match_document_to_clauses(
                    file_path, 
                    clauses, 
                    threshold=match_threshold
                )
                
                # Take top N matches
                top_matches = matches[:max_matches_per_doc]
                
                if top_matches:
                    results['documents_matched'] += 1
                    results['total_matches'] += len(top_matches)
                    results['matches'].append({
                        'file_path': file_path,
                        'file_name': filename,
                        'clause_matches': top_matches
                    })
                else:
                    results['unmatched_documents'].append({
                        'file_path': file_path,
                        'file_name': filename
                    })
        
        return results
    
    def calculate_file_hash(self, file_path: str) -> str:
        """Calculate MD5 hash of file for change detection."""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()


# ==================== USAGE EXAMPLE ====================

if __name__ == "__main__":
    # Example usage
    scanner = EnhancedDocumentScanner()
    
    # Example clauses (would come from database)
    example_clauses = [
        {
            'id': 1,
            'clause_number': '6.1.2',
            'title': 'Hazard Identification and Assessment of Risks',
            'description': 'Process for ongoing hazard identification, risk assessment and determination of necessary controls'
        },
        {
            'id': 2,
            'clause_number': '5.2',
            'title': 'OH&S Policy',
            'description': 'Top management shall establish, implement and maintain an OH&S policy'
        }
    ]
    
    # Scan a folder
    results = scanner.scan_folder(
        folder_path='/path/to/documents',
        clauses=example_clauses,
        match_threshold=0.3,
        max_matches_per_doc=3
    )
    
    print(f"Documents scanned: {results['documents_scanned']}")
    print(f"Documents matched: {results['documents_matched']}")
    print(f"Total matches: {results['total_matches']}")
    print(f"\nTop matches:")
    for match in results['matches'][:5]:
        print(f"\n{match['file_name']}:")
        for clause_id, score, reason in match['clause_matches']:
            print(f"  - Clause {clause_id}: {score:.2f} - {reason}")